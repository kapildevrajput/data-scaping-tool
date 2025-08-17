import streamlit as st
import pandas as pd
import json
import requests
from io import BytesIO
from docx import Document
from fpdf import FPDF

# -------------------
# Helper Functions
# -------------------

def fetch_api_data(url, method="GET", headers=None, json_body=None):
    try:
        if method.upper() == "GET":
            response = requests.get(url, headers=headers)
        else:
            response = requests.post(url, headers=headers, json=json_body)
        response.raise_for_status()
        data = response.json()
        # Convert JSON to DataFrame
        if isinstance(data, list):
            df = pd.DataFrame(data)
        elif isinstance(data, dict):
            df = pd.DataFrame([data])
        else:
            df = pd.DataFrame()
        return df
    except Exception as e:
        st.error(f"Error fetching API: {e}")
        return pd.DataFrame()

def export_to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

def export_to_excel(df):
    buffer = BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)
    return buffer

def export_to_word(df):
    doc = Document()
    doc.add_heading('Data Export', 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_pdf(df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Header
    pdf.cell(200, 10, txt="Data Export", ln=True, align='C')
    pdf.ln(10)
    # Table
    col_width = pdf.w / (len(df.columns) + 1)
    row_height = pdf.font_size * 1.5
    for col in df.columns:
        pdf.cell(col_width, row_height, txt=str(col), border=1)
    pdf.ln(row_height)
    for _, row in df.iterrows():
        for item in row:
            pdf.cell(col_width, row_height, txt=str(item), border=1)
        pdf.ln(row_height)
    
    # Return PDF as bytes
    return pdf.output(dest='S').encode('latin1')


# -------------------
# Streamlit UI
# -------------------

st.title("📊 Data Scraping Tool")

tab1, tab2 = st.tabs(["API Fetch", "Upload File"])

df = pd.DataFrame()

with tab1:
    st.header("Fetch Data from API")
    url = st.text_input("API URL")
    method = st.selectbox("Method", ["GET", "POST"])


    st.subheader("Headers (Optional)")
    headers_input = st.text_area("JSON format", value='{}', key="headers_input")

    try:
        headers = json.loads(headers_input)
    except:
        headers = None
    
    st.subheader("JSON Body (for POST)")
    json_input = st.text_area("JSON format", value='{}', key="json_body_input")
    try:
        json_body = json.loads(json_input)
    except:
        json_body = None
    
    if st.button("Fetch Data"):
        if url:
            df = fetch_api_data(url, method, headers, json_body)
            st.dataframe(df)
        else:
            st.warning("Please enter an API URL.")

with tab2:  # Or create a new tab "Upload File"
    st.header("Upload CSV, Excel or JSON")

    uploaded_file = st.file_uploader(
        "Choose a file",
        type=["csv", "xlsx", "json"]
    )

    json_path = st.text_input(
        "Or enter path to a local JSON file (optional)"
    )

    if uploaded_file:
        try:
            if uploaded_file.name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)

            elif uploaded_file.name.endswith(".xlsx"):
                df = pd.read_excel(uploaded_file)

            elif uploaded_file.name.endswith(".json"):
                # Try normal JSON first
                json_str = uploaded_file.read().decode("utf-8")
                try:
                    df = pd.json_normalize(json.loads(json_str))
                except json.JSONDecodeError:
                    # If normal JSON fails, assume it's JSON Lines
                    uploaded_file.seek(0)
                    df = pd.read_json(uploaded_file, lines=True)

            st.dataframe(df)

        except Exception as e:
            st.error(f"Error reading uploaded file: {e}")

    elif json_path:
        try:
            try:
                # Try normal JSON
                with open(json_path, "r", encoding="utf-8") as f:
                    json_str = f.read()
                df = pd.json_normalize(json.loads(json_str))
            except json.JSONDecodeError:
                # Try JSON Lines
                df = pd.read_json(json_path, lines=True)

            st.dataframe(df)

        except Exception as e:
            st.error(f"Error reading JSON file: {e}")
# -------------------           
# Temporary code for tab 3
# -------------------




# -------------------
# Export Section
# -------------------

if not df.empty:
    st.header("Export Data")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        csv_data = export_to_csv(df)
        st.download_button("Download CSV", csv_data, file_name="data.csv", mime="text/csv")
    with col2:
        excel_data = export_to_excel(df)
        st.download_button("Download Excel", excel_data, file_name="data.xlsx", mime="application/vnd.ms-excel")
    with col3:
        word_data = export_to_word(df)
        st.download_button("Download Word", word_data, file_name="data.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col4:
        pdf_data = export_to_pdf(df)
        st.download_button("Download PDF", pdf_data, file_name="data.pdf", mime="application/pdf")
